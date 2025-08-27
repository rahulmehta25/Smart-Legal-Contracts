"""
Document Processing Service

Comprehensive document processor supporting multiple formats:
- PDF files with text extraction and OCR
- Microsoft Word documents (DOCX)
- Plain text files (TXT)  
- HTML documents with structure preservation
- Rich text format (RTF)
- OpenDocument Text (ODT)

Features:
- Automatic file type detection
- Structure detection (headings, paragraphs, lists)
- Table extraction and formatting
- Image handling and extraction
- Metadata preservation
- Error handling and validation
"""

import io
import logging
import mimetypes
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Union, Any, Tuple
from dataclasses import dataclass, field
from enum import Enum
import hashlib

# Document processing libraries
import docx
import docx2txt
from bs4 import BeautifulSoup, NavigableString
import html2text
import filetype
import magic
import chardet
import zipfile

# Import our PDF service
from .pdf_service import PDFProcessor, ExtractionResult as PDFExtractionResult

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Supported document types"""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    HTML = "html"
    HTM = "htm"
    RTF = "rtf"
    ODT = "odt"
    UNKNOWN = "unknown"


class StructureType(Enum):
    """Document structure elements"""
    PARAGRAPH = "paragraph"
    HEADING = "heading"
    LIST = "list"
    TABLE = "table"
    IMAGE = "image"
    LINK = "link"
    FOOTER = "footer"
    HEADER = "header"


@dataclass
class DocumentElement:
    """A single document structure element"""
    type: StructureType
    content: str
    level: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)
    position: int = 0


@dataclass
class TableData:
    """Structured table data"""
    headers: List[str]
    rows: List[List[str]]
    caption: Optional[str] = None
    position: int = 0


@dataclass
class ImageData:
    """Image metadata and content"""
    filename: str
    description: Optional[str]
    size: Tuple[int, int]
    format: str
    content: Optional[bytes] = None
    position: int = 0


@dataclass
class DocumentMetadata:
    """Document metadata"""
    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    keywords: Optional[str] = None
    created: Optional[str] = None
    modified: Optional[str] = None
    language: Optional[str] = None
    page_count: int = 0
    word_count: int = 0
    character_count: int = 0
    file_size: int = 0
    file_type: DocumentType = DocumentType.UNKNOWN
    encoding: Optional[str] = None
    checksum: Optional[str] = None


@dataclass 
class ProcessingResult:
    """Complete document processing result"""
    text: str
    structured_content: List[DocumentElement]
    tables: List[TableData]
    images: List[ImageData]
    metadata: DocumentMetadata
    processing_time: float
    extraction_method: str
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)


class DocumentProcessor:
    """Multi-format document processor"""
    
    def __init__(self, 
                 extract_images: bool = True,
                 preserve_formatting: bool = True,
                 detect_structure: bool = True):
        """
        Initialize document processor
        
        Args:
            extract_images: Whether to extract images from documents
            preserve_formatting: Whether to preserve document formatting
            detect_structure: Whether to detect document structure
        """
        self.extract_images = extract_images
        self.preserve_formatting = preserve_formatting 
        self.detect_structure = detect_structure
        self.pdf_processor = PDFProcessor()
        
        # HTML to text converter
        self.html_converter = html2text.HTML2Text()
        self.html_converter.ignore_links = False
        self.html_converter.ignore_images = False
        self.html_converter.body_width = 0  # No line wrapping
    
    def process_document(self, 
                        file_path: Union[str, Path, bytes],
                        filename: Optional[str] = None,
                        password: Optional[str] = None) -> ProcessingResult:
        """
        Process document of any supported type
        
        Args:
            file_path: Path to file or file bytes
            filename: Original filename (needed for bytes input)
            password: Password for encrypted documents
            
        Returns:
            ProcessingResult with extracted content and metadata
        """
        import time
        start_time = time.time()
        
        try:
            # Detect file type
            doc_type = self._detect_file_type(file_path, filename)
            
            # Calculate file hash for deduplication
            if isinstance(file_path, (str, Path)):
                with open(file_path, 'rb') as f:
                    file_data = f.read()
            else:
                file_data = file_path
            
            checksum = hashlib.md5(file_data).hexdigest()
            
            # Process based on file type
            if doc_type == DocumentType.PDF:
                result = self._process_pdf(file_path, password)
            elif doc_type == DocumentType.DOCX:
                result = self._process_docx(file_path, password)
            elif doc_type in [DocumentType.TXT]:
                result = self._process_text(file_path)
            elif doc_type in [DocumentType.HTML, DocumentType.HTM]:
                result = self._process_html(file_path)
            elif doc_type == DocumentType.RTF:
                result = self._process_rtf(file_path)
            elif doc_type == DocumentType.ODT:
                result = self._process_odt(file_path)
            else:
                # Try to process as text if unknown
                try:
                    result = self._process_text(file_path)
                    result.warnings.append(f"Unknown file type {doc_type}, processed as text")
                except Exception as e:
                    logger.error(f"Failed to process unknown file type: {e}")
                    raise ValueError(f"Unsupported file type: {doc_type}")
            
            # Update metadata
            result.metadata.file_size = len(file_data)
            result.metadata.file_type = doc_type
            result.metadata.checksum = checksum
            result.processing_time = time.time() - start_time
            
            return result
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return ProcessingResult(
                text="",
                structured_content=[],
                tables=[],
                images=[],
                metadata=DocumentMetadata(
                    file_type=DocumentType.UNKNOWN,
                    file_size=len(file_data) if 'file_data' in locals() else 0
                ),
                processing_time=time.time() - start_time,
                extraction_method="failed",
                errors=[str(e)]
            )
    
    def _detect_file_type(self, 
                         file_path: Union[str, Path, bytes], 
                         filename: Optional[str] = None) -> DocumentType:
        """Detect document file type"""
        
        try:
            # Try filetype library first (most reliable)
            if isinstance(file_path, (str, Path)):
                kind = filetype.guess(str(file_path))
                if kind:
                    mime = kind.mime
                else:
                    # Fallback to python-magic
                    mime = magic.from_file(str(file_path), mime=True)
            else:
                kind = filetype.guess(file_path)
                if kind:
                    mime = kind.mime
                else:
                    # Fallback to magic from bytes
                    mime = magic.from_buffer(file_path, mime=True)
            
            # Map MIME types to document types
            mime_map = {
                'application/pdf': DocumentType.PDF,
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocumentType.DOCX,
                'application/msword': DocumentType.DOC,
                'text/plain': DocumentType.TXT,
                'text/html': DocumentType.HTML,
                'application/rtf': DocumentType.RTF,
                'text/rtf': DocumentType.RTF,
                'application/vnd.oasis.opendocument.text': DocumentType.ODT,
            }
            
            if mime in mime_map:
                return mime_map[mime]
            
            # Fallback to file extension
            if filename:
                ext = Path(filename).suffix.lower()
            elif isinstance(file_path, (str, Path)):
                ext = Path(file_path).suffix.lower()
            else:
                ext = ""
            
            ext_map = {
                '.pdf': DocumentType.PDF,
                '.docx': DocumentType.DOCX,
                '.doc': DocumentType.DOC,
                '.txt': DocumentType.TXT,
                '.html': DocumentType.HTML,
                '.htm': DocumentType.HTM,
                '.rtf': DocumentType.RTF,
                '.odt': DocumentType.ODT,
            }
            
            return ext_map.get(ext, DocumentType.UNKNOWN)
            
        except Exception as e:
            logger.warning(f"File type detection failed: {e}")
            return DocumentType.UNKNOWN
    
    def _process_pdf(self, 
                    file_path: Union[str, Path, bytes], 
                    password: Optional[str] = None) -> ProcessingResult:
        """Process PDF document"""
        
        pdf_result = self.pdf_processor.extract_text(file_path, password, 
                                                   preserve_layout=self.preserve_formatting,
                                                   extract_images=self.extract_images)
        
        structured_content = []
        tables = []
        images = []
        
        # Convert PDF pages to structured content
        for page in pdf_result.pages:
            if page.get('text'):
                # Simple paragraph detection
                paragraphs = [p.strip() for p in page['text'].split('\n\n') if p.strip()]
                for i, para in enumerate(paragraphs):
                    structured_content.append(DocumentElement(
                        type=StructureType.PARAGRAPH,
                        content=para,
                        position=len(structured_content),
                        metadata={'page': page['page_number']}
                    ))
            
            # Add table info if available
            if page.get('tables', 0) > 0:
                tables.append(TableData(
                    headers=[],
                    rows=[],
                    caption=f"Table from page {page['page_number']}",
                    position=page['page_number']
                ))
        
        # Convert PDF metadata
        pdf_meta = pdf_result.metadata
        metadata = DocumentMetadata(
            title=pdf_meta.title,
            author=pdf_meta.author,
            subject=pdf_meta.subject,
            created=pdf_meta.creation_date,
            modified=pdf_meta.modification_date,
            page_count=pdf_meta.pages,
            word_count=len(pdf_result.text.split()),
            character_count=len(pdf_result.text),
            language=pdf_meta.language,
            encoding="binary"
        )
        
        return ProcessingResult(
            text=pdf_result.text,
            structured_content=structured_content,
            tables=tables,
            images=images,
            metadata=metadata,
            processing_time=pdf_result.processing_time,
            extraction_method=pdf_result.extraction_method,
            errors=pdf_result.errors,
            warnings=pdf_result.warnings
        )
    
    def _process_docx(self, 
                     file_path: Union[str, Path, bytes], 
                     password: Optional[str] = None) -> ProcessingResult:
        """Process Microsoft Word DOCX document"""
        
        structured_content = []
        tables = []
        images = []
        
        try:
            # Handle both file paths and bytes
            if isinstance(file_path, (str, Path)):
                doc = docx.Document(file_path)
                # Also extract simple text
                simple_text = docx2txt.process(str(file_path))
            else:
                doc = docx.Document(io.BytesIO(file_path))
                # Save to temp file for docx2txt
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                    tmp.write(file_path)
                    tmp_path = tmp.name
                try:
                    simple_text = docx2txt.process(tmp_path)
                finally:
                    Path(tmp_path).unlink(missing_ok=True)
            
            # Extract structured content
            position = 0
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    para = None
                    for p in doc.paragraphs:
                        if p._element == element:
                            para = p
                            break
                    
                    if para and para.text.strip():
                        # Detect heading level
                        level = 0
                        if para.style.name.startswith('Heading'):
                            try:
                                level = int(para.style.name.split()[-1])
                                elem_type = StructureType.HEADING
                            except:
                                elem_type = StructureType.PARAGRAPH
                        else:
                            elem_type = StructureType.PARAGRAPH
                        
                        structured_content.append(DocumentElement(
                            type=elem_type,
                            content=para.text,
                            level=level,
                            position=position,
                            metadata={'style': para.style.name}
                        ))
                        position += 1
                
                elif element.tag.endswith('tbl'):  # Table
                    table = None
                    for t in doc.tables:
                        if t._element == element:
                            table = t
                            break
                    
                    if table:
                        headers = []
                        rows = []
                        
                        for i, row in enumerate(table.rows):
                            row_cells = [cell.text.strip() for cell in row.cells]
                            if i == 0 and self._looks_like_header(row_cells):
                                headers = row_cells
                            else:
                                rows.append(row_cells)
                        
                        tables.append(TableData(
                            headers=headers,
                            rows=rows,
                            position=position
                        ))
                        position += 1
            
            # Extract images if requested
            if self.extract_images:
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            image_part = rel.target_part
                            images.append(ImageData(
                                filename=rel.target_ref.split('/')[-1],
                                description=None,
                                size=(0, 0),  # Would need PIL to get actual size
                                format=rel.target_ref.split('.')[-1],
                                content=image_part.blob,
                                position=len(images)
                            ))
                        except Exception as e:
                            logger.warning(f"Failed to extract image: {e}")
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = DocumentMetadata(
                title=core_props.title,
                author=core_props.author,
                subject=core_props.subject,
                keywords=core_props.keywords,
                created=str(core_props.created) if core_props.created else None,
                modified=str(core_props.modified) if core_props.modified else None,
                word_count=len(simple_text.split()),
                character_count=len(simple_text),
                encoding="utf-8"
            )
            
            return ProcessingResult(
                text=simple_text,
                structured_content=structured_content,
                tables=tables,
                images=images,
                metadata=metadata,
                processing_time=0,
                extraction_method="docx_native",
                errors=[],
                warnings=[]
            )
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            raise
    
    def _process_text(self, file_path: Union[str, Path, bytes]) -> ProcessingResult:
        """Process plain text file"""
        
        try:
            # Read and detect encoding
            if isinstance(file_path, (str, Path)):
                with open(file_path, 'rb') as f:
                    raw_data = f.read()
            else:
                raw_data = file_path
            
            # Detect encoding
            encoding_result = chardet.detect(raw_data)
            encoding = encoding_result['encoding'] or 'utf-8'
            confidence = encoding_result['confidence'] or 0
            
            # Decode text
            try:
                text = raw_data.decode(encoding)
            except UnicodeDecodeError:
                # Fallback to utf-8 with error handling
                text = raw_data.decode('utf-8', errors='replace')
                encoding = 'utf-8'
            
            structured_content = []
            
            # Simple structure detection
            if self.detect_structure:
                lines = text.split('\n')
                position = 0
                
                current_paragraph = []
                for line in lines:
                    line = line.strip()
                    
                    if not line:  # Empty line - end of paragraph
                        if current_paragraph:
                            para_text = ' '.join(current_paragraph)
                            structured_content.append(DocumentElement(
                                type=StructureType.PARAGRAPH,
                                content=para_text,
                                position=position
                            ))
                            position += 1
                            current_paragraph = []
                    else:
                        current_paragraph.append(line)
                
                # Add final paragraph
                if current_paragraph:
                    para_text = ' '.join(current_paragraph)
                    structured_content.append(DocumentElement(
                        type=StructureType.PARAGRAPH,
                        content=para_text,
                        position=position
                    ))
            
            metadata = DocumentMetadata(
                word_count=len(text.split()),
                character_count=len(text),
                encoding=encoding
            )
            
            warnings = []
            if confidence < 0.8:
                warnings.append(f"Low encoding confidence: {confidence:.2f}")
            
            return ProcessingResult(
                text=text,
                structured_content=structured_content,
                tables=[],
                images=[],
                metadata=metadata,
                processing_time=0,
                extraction_method="text_native",
                errors=[],
                warnings=warnings
            )
            
        except Exception as e:
            logger.error(f"Text processing failed: {e}")
            raise
    
    def _process_html(self, file_path: Union[str, Path, bytes]) -> ProcessingResult:
        """Process HTML document"""
        
        try:
            # Read HTML content
            if isinstance(file_path, (str, Path)):
                with open(file_path, 'rb') as f:
                    raw_data = f.read()
            else:
                raw_data = file_path
            
            # Detect encoding
            encoding_result = chardet.detect(raw_data)
            encoding = encoding_result['encoding'] or 'utf-8'
            
            try:
                html_content = raw_data.decode(encoding)
            except UnicodeDecodeError:
                html_content = raw_data.decode('utf-8', errors='replace')
                encoding = 'utf-8'
            
            # Parse with BeautifulSoup
            soup = BeautifulSoup(html_content, 'lxml')
            
            # Convert to plain text
            text = self.html_converter.handle(html_content)
            
            structured_content = []
            tables = []
            images = []
            position = 0
            
            # Extract structured content
            if self.detect_structure:
                # Extract headings
                for i in range(1, 7):
                    for heading in soup.find_all(f'h{i}'):
                        if heading.text.strip():
                            structured_content.append(DocumentElement(
                                type=StructureType.HEADING,
                                content=heading.text.strip(),
                                level=i,
                                position=position,
                                metadata={'tag': f'h{i}'}
                            ))
                            position += 1
                
                # Extract paragraphs
                for para in soup.find_all('p'):
                    if para.text.strip():
                        structured_content.append(DocumentElement(
                            type=StructureType.PARAGRAPH,
                            content=para.text.strip(),
                            position=position
                        ))
                        position += 1
                
                # Extract lists
                for ul in soup.find_all(['ul', 'ol']):
                    list_items = [li.text.strip() for li in ul.find_all('li') if li.text.strip()]
                    if list_items:
                        structured_content.append(DocumentElement(
                            type=StructureType.LIST,
                            content='\n'.join(f"• {item}" for item in list_items),
                            position=position,
                            metadata={'tag': ul.name, 'items': len(list_items)}
                        ))
                        position += 1
                
                # Extract tables
                for table in soup.find_all('table'):
                    headers = []
                    rows = []
                    
                    # Try to find header row
                    header_row = table.find('thead')
                    if header_row:
                        headers = [th.text.strip() for th in header_row.find_all(['th', 'td'])]
                    
                    # Extract data rows
                    tbody = table.find('tbody') or table
                    for row in tbody.find_all('tr'):
                        if not headers and row == tbody.find('tr'):
                            # First row might be headers
                            row_cells = [cell.text.strip() for cell in row.find_all(['th', 'td'])]
                            if self._looks_like_header(row_cells):
                                headers = row_cells
                                continue
                        
                        row_cells = [cell.text.strip() for cell in row.find_all(['th', 'td'])]
                        if any(cell for cell in row_cells):  # Skip empty rows
                            rows.append(row_cells)
                    
                    if rows or headers:
                        tables.append(TableData(
                            headers=headers,
                            rows=rows,
                            position=position
                        ))
                        position += 1
                
                # Extract images
                if self.extract_images:
                    for img in soup.find_all('img'):
                        src = img.get('src')
                        alt = img.get('alt', '')
                        if src:
                            images.append(ImageData(
                                filename=src.split('/')[-1] if '/' in src else src,
                                description=alt,
                                size=(0, 0),
                                format=src.split('.')[-1] if '.' in src else 'unknown',
                                position=len(images)
                            ))
            
            # Extract metadata
            title = soup.find('title')
            title_text = title.text.strip() if title else None
            
            # Look for meta tags
            meta_author = soup.find('meta', {'name': 'author'})
            author = meta_author.get('content') if meta_author else None
            
            meta_description = soup.find('meta', {'name': 'description'})
            description = meta_description.get('content') if meta_description else None
            
            metadata = DocumentMetadata(
                title=title_text,
                author=author,
                subject=description,
                word_count=len(text.split()),
                character_count=len(text),
                encoding=encoding
            )
            
            return ProcessingResult(
                text=text,
                structured_content=structured_content,
                tables=tables,
                images=images,
                metadata=metadata,
                processing_time=0,
                extraction_method="html_bs4",
                errors=[],
                warnings=[]
            )
            
        except Exception as e:
            logger.error(f"HTML processing failed: {e}")
            raise
    
    def _process_rtf(self, file_path: Union[str, Path, bytes]) -> ProcessingResult:
        """Process RTF document (basic implementation)"""
        
        # For now, treat as text and strip RTF codes
        # A full RTF parser would be more complex
        
        try:
            if isinstance(file_path, (str, Path)):
                with open(file_path, 'rb') as f:
                    raw_data = f.read()
            else:
                raw_data = file_path
            
            # Try to decode as text
            try:
                rtf_content = raw_data.decode('utf-8')
            except UnicodeDecodeError:
                rtf_content = raw_data.decode('latin1', errors='replace')
            
            # Basic RTF cleanup - remove control codes
            import re
            
            # Remove RTF header
            text = re.sub(r'^{\\rtf.*?}', '', rtf_content)
            
            # Remove control sequences
            text = re.sub(r'\\[a-z]+\d*', '', text)
            text = re.sub(r'[{}\\]', '', text)
            
            # Clean up whitespace
            text = re.sub(r'\s+', ' ', text).strip()
            
            structured_content = []
            if text.strip():
                structured_content.append(DocumentElement(
                    type=StructureType.PARAGRAPH,
                    content=text,
                    position=0
                ))
            
            metadata = DocumentMetadata(
                word_count=len(text.split()),
                character_count=len(text),
                encoding="rtf"
            )
            
            return ProcessingResult(
                text=text,
                structured_content=structured_content,
                tables=[],
                images=[],
                metadata=metadata,
                processing_time=0,
                extraction_method="rtf_basic",
                errors=[],
                warnings=["RTF processing is basic - some formatting may be lost"]
            )
            
        except Exception as e:
            logger.error(f"RTF processing failed: {e}")
            raise
    
    def _process_odt(self, file_path: Union[str, Path, bytes]) -> ProcessingResult:
        """Process OpenDocument Text document"""
        
        try:
            # ODT files are ZIP archives
            if isinstance(file_path, (str, Path)):
                with zipfile.ZipFile(file_path, 'r') as odt_zip:
                    content_xml = odt_zip.read('content.xml')
            else:
                with zipfile.ZipFile(io.BytesIO(file_path), 'r') as odt_zip:
                    content_xml = odt_zip.read('content.xml')
            
            # Parse the content.xml
            soup = BeautifulSoup(content_xml, 'xml')
            
            # Extract text content
            text_elements = soup.find_all(['text:p', 'text:h'])
            text_parts = []
            structured_content = []
            position = 0
            
            for elem in text_elements:
                text_content = elem.get_text().strip()
                if text_content:
                    text_parts.append(text_content)
                    
                    # Determine element type
                    if elem.name == 'text:h':
                        elem_type = StructureType.HEADING
                        level = int(elem.get('text:outline-level', 1))
                    else:
                        elem_type = StructureType.PARAGRAPH
                        level = 0
                    
                    structured_content.append(DocumentElement(
                        type=elem_type,
                        content=text_content,
                        level=level,
                        position=position
                    ))
                    position += 1
            
            text = '\n\n'.join(text_parts)
            
            # Try to extract metadata from meta.xml
            metadata = DocumentMetadata(
                word_count=len(text.split()),
                character_count=len(text),
                encoding="xml"
            )
            
            try:
                if isinstance(file_path, (str, Path)):
                    with zipfile.ZipFile(file_path, 'r') as odt_zip:
                        meta_xml = odt_zip.read('meta.xml')
                else:
                    with zipfile.ZipFile(io.BytesIO(file_path), 'r') as odt_zip:
                        meta_xml = odt_zip.read('meta.xml')
                
                meta_soup = BeautifulSoup(meta_xml, 'xml')
                title_elem = meta_soup.find('dc:title')
                author_elem = meta_soup.find('dc:creator')
                subject_elem = meta_soup.find('dc:subject')
                
                if title_elem:
                    metadata.title = title_elem.get_text()
                if author_elem:
                    metadata.author = author_elem.get_text()
                if subject_elem:
                    metadata.subject = subject_elem.get_text()
                    
            except Exception as e:
                logger.warning(f"Failed to extract ODT metadata: {e}")
            
            return ProcessingResult(
                text=text,
                structured_content=structured_content,
                tables=[],
                images=[],
                metadata=metadata,
                processing_time=0,
                extraction_method="odt_xml",
                errors=[],
                warnings=[]
            )
            
        except Exception as e:
            logger.error(f"ODT processing failed: {e}")
            raise
    
    def _looks_like_header(self, row_cells: List[str]) -> bool:
        """Determine if a table row looks like a header row"""
        
        if not row_cells:
            return False
        
        # Simple heuristics
        non_empty = [cell for cell in row_cells if cell.strip()]
        if len(non_empty) < len(row_cells) / 2:  # Too many empty cells
            return False
        
        # Check if cells are short (typical for headers)
        avg_length = sum(len(cell) for cell in non_empty) / len(non_empty)
        if avg_length > 50:  # Headers are usually short
            return False
        
        return True


# Convenience function
def process_document(file_path: Union[str, Path, bytes], 
                    filename: Optional[str] = None,
                    password: Optional[str] = None,
                    extract_images: bool = True) -> ProcessingResult:
    """
    Process document of any supported type
    
    Args:
        file_path: Path to file or file bytes
        filename: Original filename (needed for bytes input)  
        password: Password for encrypted documents
        extract_images: Whether to extract images
        
    Returns:
        ProcessingResult with extracted content and metadata
    """
    processor = DocumentProcessor(extract_images=extract_images)
    return processor.process_document(file_path, filename, password)